# src/core/memo_generator.py
from typing import Dict, List, Optional
import json
from src.core.consensus_chain import ConsensusChain
from src.core.fallback_handler import FallbackHandler
from src.utils.config import Config


class MemoGenerator:
    """备忘录生成器(三层架构)"""

    def __init__(self, consensus_chain: ConsensusChain, llm_client=None,
                 client_profile: Dict = None, fallback_handler: FallbackHandler = None):
        self.consensus_chain = consensus_chain
        self.llm_client = llm_client
        self.client_profile = client_profile or {}  # 客户档案（从外部注入）
        self.fallback_handler = fallback_handler or FallbackHandler()

    def _strip_metadata(self, direction: Dict) -> Dict:
        """剥离内部元数据字段（来源等），防止写入Word或传给LLM"""
        return {
            "方向": direction.get("方向", "")
            # 注意：不包含"来源"字段
        }

    def extract_data(self) -> Dict:
        """第一层: 数据提取(确定性规则)"""
        data = {
            "facts": [],
            "consensus": [],
            "pending": [],
            "client_profile": self.client_profile  # 使用注入的客户档案
        }

        # 提取已确认事实
        for record in self.consensus_chain.get_confirmed_facts():
            data["facts"].append({
                "stage": record.stage,
                "content": record.content,
                "source": record.source
            })

        # 提取已确认判断
        for record in self.consensus_chain.get_confirmed_consensus():
            data["consensus"].append({
                "content": record.content,
                "source": record.source,
                "recommendation": record.recommendation
            })

        # 提取待确认判断
        for record in self.consensus_chain.get_pending_consensus():
            data["pending"].append({
                "content": record.content
            })

        return data

    def generate_structure(self) -> Dict:
        """第二层: 结构组装(模板+规则)

        设计文档 4.5 节和 7.3 节映射表：
        - 一、问题重构
        - 二、关键发现（按优先级排序）
        - 三、初步建议方向
        - 四、需要进一步访谈
        - 五、建议下一步合作方式（服务包推荐）
        """
        data = self.extract_data()

        structure = {
            "chapters": {}
        }

        # 一、问题重构
        structure["chapters"]["问题重构"] = {
            "原始诉求": data["client_profile"].get("显性诉求", ""),
            "核心问题": data["consensus"][0]["content"] if data["consensus"] else ""
        }

        # 二、关键发现（按优先级排序：source=candidate_selected > 时间戳最早 > 时间戳最新）
        strategy_facts = self._sort_facts_by_priority([
            f for f in data["facts"] if f["stage"] == "战略梳理"
        ])
        business_facts = self._sort_facts_by_priority([
            f for f in data["facts"] if f["stage"] == "商业模式"
        ])

        structure["chapters"]["关键发现"] = {
            "战略层面": [f["content"] for f in strategy_facts[:3]],
            "商业模式层面": [f["content"] for f in business_facts[:3]]
        }

        # 超过3条的降级到"需要进一步访谈"
        extra_facts = strategy_facts[3:] + business_facts[3:]
        structure["chapters"]["_extra_facts"] = [f["content"] for f in extra_facts]

        # 三、初步建议方向
        structure["chapters"]["初步建议方向"] = [
            {
                "方向": c["recommendation"] or c["content"],
                "来源": "系统生成" if c["source"] == "candidate_selected" else "顾问判断"
            }
            for c in data["consensus"]
        ]

        # 四、需要进一步访谈
        structure["chapters"]["需要进一步访谈"] = [
            p["content"] for p in data["pending"]
        ] + structure["chapters"].pop("_extra_facts", [])

        # 五、建议下一步合作方式（设计文档 7.6 节）
        structure["chapters"]["建议下一步合作方式"] = self._generate_service_recommendation(data)

        return structure

    def _sort_facts_by_priority(self, facts: List[Dict]) -> List[Dict]:
        """按优先级排序事实

        规则（设计文档 7.3 节）：
        1. source=candidate_selected 优先
        2. 时间戳最早优先
        3. 时间戳最新优先
        """
        def priority_key(fact):
            # source=candidate_selected 排最前
            source_priority = 0 if fact.get("source") == "candidate_selected" else 1
            return source_priority

        return sorted(facts, key=priority_key)

    def _generate_service_recommendation(self, data: Dict) -> Dict:
        """生成服务包推荐（设计文档 7.6 节）

        逻辑：
        - 共识链确认条数 >= 5 且完整度 >= 60%：推荐"深度诊断服务包"
        - 共识链确认条数 >= 3 且完整度 >= 40%：推荐"初步诊断服务包"
        - 其他：推荐"免费初步沟通"
        """
        confirmed_count = len(data["consensus"]) + len(data["facts"])
        completeness = self._calc_profile_completeness(data["client_profile"])

        if confirmed_count >= 5 and completeness >= 0.6:
            return {
                "推荐服务包": "深度诊断服务包",
                "理由": f"已确认{confirmed_count}条共识，客户档案完整度{completeness:.0%}，建议进入深度诊断阶段",
                "下一步": "安排2-3次深度访谈，聚焦关键决策点"
            }
        elif confirmed_count >= 3 and completeness >= 0.4:
            return {
                "推荐服务包": "初步诊断服务包",
                "理由": f"已确认{confirmed_count}条共识，建议进一步明确诊断方向",
                "下一步": "补充关键信息，完善客户档案"
            }
        else:
            return {
                "推荐服务包": "免费初步沟通",
                "理由": f"当前共识条数{confirmed_count}条，建议继续建立信任关系",
                "下一步": "聚焦客户痛点，收集更多背景信息"
            }

    def _calc_profile_completeness(self, profile: Dict) -> float:
        """计算客户档案完整度"""
        if not profile:
            return 0.0

        required_fields = [
            "产品线", "客户群体", "收入结构", "毛利结构",
            "交付情况", "资源分布", "战略目标", "显性诉求"
        ]

        filled = sum(1 for f in required_fields if profile.get(f) and len(str(profile.get(f, ""))) >= 5)
        return filled / len(required_fields)

    def generate_word(self, output_path: str):
        """生成Word文档(第三层: AI润色在Day 2实现)"""
        from docx import Document

        structure = self.generate_structure()
        doc = Document()

        # 标题
        doc.add_heading('客户初步诊断备忘录', 0)

        # 一、问题重构
        doc.add_heading('一、问题重构', level=1)
        problem = structure["chapters"]["问题重构"]
        doc.add_paragraph(f"原始诉求: {problem['原始诉求']}")
        doc.add_paragraph(f"诊断后的核心问题: {problem['核心问题']}")

        # 二、关键发现
        doc.add_heading('二、关键发现', level=1)
        findings = structure["chapters"]["关键发现"]
        for stage, facts in findings.items():
            if facts:
                doc.add_heading(stage, level=2)
                for fact in facts:
                    doc.add_paragraph(fact, style='List Bullet')

        # 三、初步建议方向
        doc.add_heading('三、初步建议方向', level=1)
        for i, direction in enumerate(structure["chapters"]["初步建议方向"], 1):
            # 剥离内部元数据后再写入Word
            clean_direction = self._strip_metadata(direction)
            doc.add_paragraph(f"方向{i}: {clean_direction['方向']}")

        # 四、需要进一步访谈
        doc.add_heading('四、需要进一步访谈', level=1)
        for item in structure["chapters"]["需要进一步访谈"]:
            doc.add_paragraph(item, style='List Bullet')

        # 五、建议下一步合作方式
        doc.add_heading('五、建议下一步合作方式', level=1)
        recommendation = structure["chapters"]["建议下一步合作方式"]
        doc.add_paragraph(f"推荐服务包: {recommendation['推荐服务包']}")
        doc.add_paragraph(f"理由: {recommendation['理由']}")
        doc.add_paragraph(f"下一步: {recommendation['下一步']}")

        doc.save(output_path)

    def polish_chapter(self, chapter_data: Dict, max_words: int = 200) -> str:
        """第三层：AI润色章节（带超时保护）

        Args:
            chapter_data: 章节数据（键值对或列表）
            max_words: 最大字数限制

        Returns:
            润色后的段落文字，或降级后的要点列表
        """
        if not self.llm_client:
            # 降级：无LLM时直接返回要点列表
            return self._format_as_bullets(chapter_data)

        prompt = f"""请将以下要点转化成连贯的段落文字。

要求：
1. 只能用下面的要点，不能添加任何额外信息
2. 字数不超过{max_words}字
3. 所有数字和专有名词必须与原始要点完全一致
4. 语气专业、客观、建设性，不使用夸张词汇

要点：
{json.dumps(chapter_data, ensure_ascii=False, indent=2)}

直接输出润色后的段落，不要其他解释。"""

        # 使用 FallbackHandler 进行超时保护
        result = self.fallback_handler.handle_llm_timeout(
            generator=lambda: self.llm_client.generate(prompt, temperature=0.3),
            timeout_seconds=Config.LLM_TIMEOUT_SECONDS,
            fallback_value=self._format_as_bullets(chapter_data)
        )

        if result.success:
            polished = result.data.get("result", "")
            # 字数截断
            if len(polished) > max_words:
                polished = polished[:max_words] + "..."
            return polished
        else:
            # 超时降级：使用原始要点
            return result.data.get("fallback_value", self._format_as_bullets(chapter_data))

    def _format_as_bullets(self, chapter_data: Dict) -> str:
        """降级：格式化为要点列表"""
        lines = []
        for key, value in chapter_data.items():
            if isinstance(value, list):
                for item in value:
                    lines.append(f"- {item}")
            else:
                lines.append(f"- {key}: {value}")
        return "\n".join(lines)
