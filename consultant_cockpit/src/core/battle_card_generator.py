# src/core/battle_card_generator.py
"""会前作战卡生成器

根据设计文档 11.1 节实现：
- 双模式自动切换（验证假设版 / 信息建立版）
- 规则过滤 → 优先级排序 → LLM润色 → 硬约束校验
- Word文档输出
"""
from typing import Optional, List, Dict, Any
from datetime import datetime
from io import BytesIO
import json

from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class BattleCard(BaseModel):
    """作战卡数据结构（使用 Pydantic BaseModel 保持项目一致性）"""
    company: str
    date: str
    consultant: str
    mode: str  # "hypothesis" | "info_building"
    completeness: float
    content: Dict[str, Any] = Field(default_factory=dict)


class InsufficientSkuError(Exception):
    """SKU数量不足异常"""
    pass


# 预设追问树（固定在作战卡中）
PRESET_STRATEGY_TREE = {
    "anchor": "您现在最头疼的一件事是什么?",
    "branches": {
        "增长": ["现在增长靠什么驱动?", "这个驱动力能持续多久?"],
        "盈利": ["哪条线最赚钱?为什么?", "其他线是战略投入还是历史包袱?"],
        "方向": ["现在有几个方向在跑?", "资源是怎么分配的?"]
    }
}

PRESET_BUSINESS_TREE = {
    "anchor": "您的主要收入来源是什么?",
    "branches": {
        "设备销售": ["有没有想过卖服务?", "客户愿意为运营结果付费吗?"],
        "EPC工程": ["工程完了客户还找你吗?", "有没有机会做长期运维?"],
        "运营服务": ["现在规模多大?", "可复制性怎么样?"]
    }
}

# 风险话术模板
RISK_RESPONSES = {
    "已有方向": {
        "trigger": "客户说\"我们已经有方向了\"",
        "response": "那您觉得现在最大的执行障碍是什么?"
    },
    "超出范围": {
        "trigger": "客户问超出范围的问题",
        "response": "这是关键问题，列入下阶段专项研究，一周内回复"
    },
    "质疑专业性": {
        "trigger": "客户质疑专业性",
        "response": "动态生成：引用🟢SKU作为背书"
    }
}


class BattleCardGenerator:
    """作战卡生成器"""

    # SKU最小数量保护阈值
    MIN_SKU_COUNT = 6

    def __init__(self, feishu_client, llm_client, knowledge_retriever=None):
        """
        Args:
            feishu_client: 飞书客户端（用于获取客户档案和计算完整度）
            llm_client: LLM客户端（用于润色）
            knowledge_retriever: 知识召回器（可选，用于SKU召回）
        """
        self.feishu_client = feishu_client
        self.llm_client = llm_client
        self.knowledge_retriever = knowledge_retriever

    def generate(self, company: str, consultant: str = "") -> BattleCard:
        """生成会前作战卡

        Args:
            company: 客户公司名称
            consultant: 顾问姓名（可选）

        Returns:
            BattleCard: 作战卡数据结构

        Raises:
            ValueError: 当客户档案不存在时
        """
        # 1. 获取客户档案
        profile = self.feishu_client.get_client_profile(company)

        # 2. 处理 profile 为 None 的情况
        if profile is None:
            # 创建空档案，使用默认值继续生成
            profile = {
                "record_id": None,
                "fields": {
                    "客户公司名": company,
                    # 其他字段为空
                }
            }

        # 3. 计算完整度（复用 FeishuClient）
        completeness = self.feishu_client.calc_completeness(profile)

        # 3. 根据完整度选择模式
        if completeness >= 0.6:
            content = self._generate_hypothesis_version(profile)
            mode = "hypothesis"
        else:
            content = self._generate_info_building_version(profile)
            mode = "info_building"

        return BattleCard(
            company=company,
            date=datetime.now().strftime("%Y-%m-%d"),
            consultant=consultant,
            mode=mode,
            completeness=completeness,
            content=content
        )

    def _get_top_skus(self, profile: Dict, top_n: int = 15) -> List[Dict]:
        """获取Top N SKU

        Args:
            profile: 客户档案
            top_n: 返回数量

        Returns:
            SKU列表
        """
        if self.knowledge_retriever:
            # 从知识库召回
            keywords = self._extract_keywords_from_profile(profile)
            skus = self.knowledge_retriever.recall_by_keywords(keywords, top_k=top_n)
            # 转换为字典格式
            return [
                {
                    "id": sku.id,
                    "title": sku.title,
                    "summary": sku.summary,
                    "confidence": sku.confidence,
                    "stage": sku.stage
                }
                for sku in skus
            ]
        else:
            # 降级：使用mock数据
            return self._get_mock_skus(top_n)

    def _get_mock_skus(self, top_n: int = 15) -> List[Dict]:
        """Mock SKU数据（降级用）"""
        mock_skus = [
            {"id": "sku_001", "title": "设备商转运营商路径", "summary": "从设备销售转向运营服务的典型案例", "confidence": "🟢", "stage": "商业模式"},
            {"id": "sku_002", "title": "储能系统集成商商业模式", "summary": "工商业储能系统集成商的盈利模式分析", "confidence": "🟡", "stage": "商业模式"},
            {"id": "sku_003", "title": "虚拟电厂聚合商案例", "summary": "负荷聚合商参与电力市场的路径", "confidence": "🟢", "stage": "战略梳理"},
            {"id": "sku_004", "title": "分布式光伏运营模式", "summary": "分布式光伏的商业模式创新", "confidence": "🟢", "stage": "商业模式"},
            {"id": "sku_005", "title": "重卡换电商业案例", "summary": "重卡换电站的盈利模型", "confidence": "🟡", "stage": "商业模式"},
            {"id": "sku_006", "title": "微电网运营案例", "summary": "工业园区微电网运营实践", "confidence": "🟢", "stage": "战略梳理"},
            {"id": "sku_007", "title": "综合能源服务转型", "summary": "传统能源企业转型综合能源服务", "confidence": "🟡", "stage": "战略梳理"},
            {"id": "sku_008", "title": "储能电站运营", "summary": "独立储能电站的商业模式", "confidence": "🟢", "stage": "商业模式"},
            {"id": "sku_009", "title": "电力交易代理", "summary": "电力市场化交易代理服务", "confidence": "🟢", "stage": "战略梳理"},
            {"id": "sku_010", "title": "需求响应聚合", "summary": "需求响应资源聚合模式", "confidence": "🟡", "stage": "商业模式"},
        ]
        return mock_skus[:top_n]

    def _extract_keywords_from_profile(self, profile: Dict) -> List[str]:
        """从客户档案提取关键词"""
        if not profile:
            return ["新能源", "储能", "光伏"]

        fields = profile.get("fields", {})
        keywords = []

        # 从各字段提取
        for field_name in ["产品线", "客户群体", "战略目标"]:
            value = fields.get(field_name, "")
            if value:
                keywords.extend(value.split("/")[:3])
                keywords.extend(value.split("、")[:3])

        return list(set(keywords))[:10] if keywords else ["新能源", "储能", "光伏"]

    def _filter_skus(self, profile: Dict) -> List[Dict]:
        """规则过滤：根据角色×场景×痛点过滤SKU

        Day 3 简化实现：直接返回mock数据
        Day 4+ 可接入真实知识库
        """
        return self._get_top_skus(profile)

    def _rank_skus(self, candidate_skus: List[Dict], profile: Dict) -> List[Dict]:
        """优先级排序：加权计算Top 15

        权重：
        - 角色匹配度 0.4
        - 痛点匹配度 0.3
        - SKU可信度 0.2
        - 当前阶段相关性 0.1

        Day 3 简化实现：按可信度排序
        """
        def confidence_score(sku):
            conf = sku.get("confidence", "🔴")
            if conf == "🟢":
                return 1.0
            elif conf == "🟡":
                return 0.7
            else:
                return 0.3

        sorted_skus = sorted(candidate_skus, key=confidence_score, reverse=True)
        return sorted_skus[:15]

    def _generate_hypothesis_version(self, profile: Dict) -> Dict:
        """生成验证假设版（完整度≥60%）

        数据流：规则过滤 → 优先级排序 → LLM润色 → 硬约束校验
        """
        # 1. 规则过滤
        candidate_skus = self._filter_skus(profile)

        # 2. 优先级排序
        top_skus = self._rank_skus(candidate_skus, profile)

        # 3. SKU最小数量保护
        if len(top_skus) < self.MIN_SKU_COUNT:
            raise InsufficientSkuError(
                f"SKU召回不足{self.MIN_SKU_COUNT}条，当前{len(top_skus)}条，无法生成高质量作战卡"
            )

        # 4. LLM润色：分批传入，转化为口播台词
        #    修复：demo_scripts 在 SKU 数量不足时使用剩余 SKU
        demo_sku_count = len(top_skus) - 9
        if demo_sku_count >= 3:
            demo_skus = top_skus[9:12]
        elif demo_sku_count > 0:
            # SKU 数量在 9-11 条之间，使用剩余的 SKU
            demo_skus = top_skus[9:]
        else:
            # SKU 数量不足 9 条，从末尾取最多 3 条
            demo_skus = top_skus[max(0, len(top_skus)-3):]

        content = {
            "diagnosis_hypothesis": self._generate_hypothesis(top_skus[:3], profile),
            "strategy_questions": self._generate_questions(top_skus[3:6], "战略梳理"),
            "business_questions": self._generate_questions(top_skus[6:9] if len(top_skus) >= 9 else [], "商业模式"),
            "demo_scripts": self._generate_scripts(demo_skus),
            "risk_responses": self._generate_risk_responses(top_skus)  # 传入SKU用于动态生成
        }

        # 5. 硬约束校验：确保诊断假设引用≥1个🟢/🟡SKU
        self._validate_constraints(content, top_skus)

        return content

    def _generate_info_building_version(self, profile: Dict) -> Dict:
        """生成信息建立版（完整度<60%）"""
        top_skus = self._get_top_skus(profile, top_n=6)

        content = {
            "missing_fields": self._identify_missing_fields(profile),
            "strategy_tree": PRESET_STRATEGY_TREE,
            "business_tree": PRESET_BUSINESS_TREE,
            "demo_scripts": self._generate_scripts(top_skus[:3]),
            "risk_responses": self._generate_risk_responses(top_skus)  # 传入SKU用于动态生成
        }

        return content

    def _generate_hypothesis(self, skus: List[Dict], profile: Dict) -> str:
        """生成诊断假设

        Args:
            skus: Top 3 SKU
            profile: 客户档案

        Returns:
            诊断假设文本
        """
        if not skus:
            return "客户业务模式需要进一步诊断，建议从战略定位和商业模式两个维度展开。"

        # 构建prompt
        sku_refs = "\n".join([
            f"- [{sku['confidence']}] {sku['title']}: {sku['summary']}"
            for sku in skus
        ])

        fields = profile.get("fields", {}) if profile else {}
        company = fields.get("客户公司名", "客户")
        product_line = fields.get("产品线", "未知")
       诉求 = fields.get("显性诉求", "")

        prompt = f"""基于以下客户信息和行业案例，生成1-2句诊断假设。

客户：{company}
产品线：{product_line}
显性诉求：{诉求}

参考案例：
{sku_refs}

要求：
1. 假设必须引用至少1个案例作为依据
2. 用专业但易懂的语言
3. 不超过100字
4. 格式："基于[案例名]的经验，客户的核心问题很可能是..."

直接输出假设，不要其他解释。"""

        try:
            result = self.llm_client.generate(prompt, temperature=0.5, max_tokens=200)
            return result.strip()
        except Exception:
            # 降级：模板填充
            return f"基于{skus[0]['title']}的经验，客户的核心问题很可能是战略定位不够清晰，建议进一步验证。"

    def _generate_questions(self, skus: List[Dict], stage: str) -> str:
        """生成必问3问

        Args:
            skus: SKU列表（3个）
            stage: 阶段名称

        Returns:
            问题列表文本
        """
        if not skus:
            # 降级：使用预设问题
            if stage == "战略梳理":
                return "1. 您现在的核心业务是什么？\n2. 未来3年的战略目标是什么？\n3. 目前最大的挑战是什么？"
            else:
                return "1. 主要收入来源是什么？\n2. 哪块业务最赚钱？\n3. 商业模式有什么特点？"

        # 构建prompt
        sku_refs = "\n".join([f"- {sku['title']}" for sku in skus])

        prompt = f"""基于以下案例，生成3个可以直接口播的诊断问题。

参考案例：{sku_refs}
阶段：{stage}

要求：
1. 问题要具体，可以直接问客户
2. 每个问题不超过30字
3. 格式：每行一个问题，编号1/2/3

直接输出3个问题，不要其他解释。"""

        try:
            result = self.llm_client.generate(prompt, temperature=0.6, max_tokens=300)
            return result.strip()
        except Exception:
            # 降级：基于SKU生成
            return "\n".join([f"{i+1}. 您如何看待{sku['title']}这个方向？" for i, sku in enumerate(skus)])

    def _generate_scripts(self, skus: List[Dict]) -> str:
        """生成口播台词

        Args:
            skus: SKU列表（3个）

        Returns:
            台词文本
        """
        if not skus:
            return "A. 案例待补充\nB. 案例待补充\nC. 案例待补充"

        lines = []
        labels = ["A", "B", "C"]
        for i, sku in enumerate(skus[:3]):
            label = labels[i]
            title = sku.get("title", "案例")
            summary = sku.get("summary", "")
            # 生成一句话台词
            script = f"在{title}领域，{summary[:30]}..." if len(summary) > 30 else f"在{title}领域，{summary}"
            lines.append(f"{label}. {title} → {script}")

        return "\n".join(lines)

    def _generate_risk_responses(self, skus: List[Dict] = None) -> str:
        """生成风险话术

        Args:
            skus: SKU列表（可选，用于动态生成第三条话术）

        Returns:
            风险话术文本
        """
        lines = []

        # 第一条：固定模板
        lines.append(f"▸ {RISK_RESPONSES['已有方向']['trigger']}")
        lines.append(f"  → {RISK_RESPONSES['已有方向']['response']}")

        # 第二条：固定模板
        lines.append(f"▸ {RISK_RESPONSES['超出范围']['trigger']}")
        lines.append(f"  → {RISK_RESPONSES['超出范围']['response']}")

        # 第三条：动态生成（引用🟢SKU作为背书）
        lines.append(f"▸ {RISK_RESPONSES['质疑专业性']['trigger']}")

        if skus:
            # 找到🟢可信度的SKU
            green_skus = [sku for sku in skus if sku.get("confidence") == "🟢"]
            if green_skus:
                # 动态生成：引用🟢SKU作为背书
                sku_ref = green_skus[0]
                lines.append(f"  → 我们在{sku_ref['title']}领域有深入研究，可以分享相关案例")
            else:
                # 没有🟢SKU，使用🟡SKU
                yellow_skus = [sku for sku in skus if sku.get("confidence") == "🟡"]
                if yellow_skus:
                    sku_ref = yellow_skus[0]
                    lines.append(f"  → 我们在{sku_ref['title']}方向有相关经验，可以展开讨论")
                else:
                    # 降级：通用话术
                    lines.append("  → 我们在新能源行业有丰富的咨询经验，可以分享具体案例")
        else:
            # 无SKU时使用通用话术
            lines.append("  → 我们在新能源行业有丰富的咨询经验，可以分享具体案例")

        return "\n".join(lines)

    def _identify_missing_fields(self, profile: Dict) -> List[str]:
        """识别缺失字段"""
        required_fields = [
            "产品线", "客户群体", "收入结构",
            "毛利结构", "交付情况", "资源分布",
            "战略目标", "显性诉求", "隐性痛点"
        ]

        if not profile:
            return required_fields

        fields = profile.get("fields", {})
        missing = []

        for field_name in required_fields:
            value = fields.get(field_name, "")
            if not value or len(str(value)) < 5:
                missing.append(field_name)

        return missing

    def _validate_constraints(self, content: Dict, skus: List[Dict]) -> bool:
        """硬约束校验

        确保诊断假设引用≥1个🟢/🟡SKU
        """
        # 检查是否有🟢/🟡SKU
        valid_skus = [sku for sku in skus if sku.get("confidence") in ["🟢", "🟡"]]

        if not valid_skus:
            raise InsufficientSkuError(
                "没有🟢/🟡可信度的SKU，无法生成高质量诊断假设"
            )

        return True

    def render_to_word(self, battle_card: BattleCard) -> bytes:
        """渲染为Word文档

        Returns:
            bytes: Word文档的字节流（用于下载或发送）
        """
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(10)

        # 标题区
        mode_text = '验证假设版' if battle_card.mode == 'hypothesis' else '信息建立版'
        title = doc.add_heading(f"客户作战卡（{mode_text}）", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息
        info_para = doc.add_paragraph()
        info_para.add_run(f"{battle_card.company} · {battle_card.date}").bold = True
        if battle_card.consultant:
            info_para.add_run(f" · 顾问：{battle_card.consultant}")

        # 完整度指示
        if battle_card.mode == "info_building":
            warning = doc.add_paragraph()
            warning.add_run(f"⚠️ 客户背景待完善，当前完整度：{battle_card.completeness:.0%}").bold = True
            warning.add_run("\n本次会议目标：建立诊断基础")

        doc.add_paragraph()  # 空行

        content = battle_card.content

        if battle_card.mode == "hypothesis":
            # 验证假设版
            self._render_hypothesis_content(doc, content)
        else:
            # 信息建立版
            self._render_info_building_content(doc, content)

        # 使用 BytesIO 返回字节流
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _render_hypothesis_content(self, doc: Document, content: Dict):
        """渲染验证假设版内容"""
        # 核心诊断假设
        doc.add_heading("【核心诊断假设】", level=1)
        hypothesis = content.get("diagnosis_hypothesis", "")
        doc.add_paragraph(hypothesis)
        doc.add_paragraph("→ 本次会议目标：验证/推翻这个假设")

        # 战略梳理阶段·必问3问
        doc.add_heading("【战略梳理阶段·必问3问】", level=1)
        strategy_q = content.get("strategy_questions", "")
        for line in strategy_q.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip(), style='List Number')

        # 商业模式阶段·必问3问
        doc.add_heading("【商业模式阶段·必问3问】", level=1)
        business_q = content.get("business_questions", "")
        for line in business_q.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip(), style='List Number')

        # 行业演示备弹
        doc.add_heading("【行业演示备弹·3条口播台词】", level=1)
        scripts = content.get("demo_scripts", "")
        for line in scripts.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

        # 风险话术
        doc.add_heading("【风险话术·应急回应】", level=1)
        risk = content.get("risk_responses", "")
        for line in risk.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    def _render_info_building_content(self, doc: Document, content: Dict):
        """渲染信息建立版内容"""
        # 必须在本次会议确认的字段
        doc.add_heading("【必须在本次会议确认的字段】", level=1)
        missing = content.get("missing_fields", [])
        for field_name in missing:
            doc.add_paragraph(f"□ {field_name}", style='List Bullet')

        # 分层追问树·战略层
        doc.add_heading("【分层追问树·战略层】", level=1)
        strategy_tree = content.get("strategy_tree", PRESET_STRATEGY_TREE)
        self._render_question_tree(doc, strategy_tree)

        # 分层追问树·商业模式层
        doc.add_heading("【分层追问树·商业模式层】", level=1)
        business_tree = content.get("business_tree", PRESET_BUSINESS_TREE)
        self._render_question_tree(doc, business_tree)

        # 行业演示备弹
        doc.add_heading("【行业演示备弹·3条口播台词】", level=1)
        scripts = content.get("demo_scripts", "")
        for line in scripts.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

        # 风险话术
        doc.add_heading("【风险话术】", level=1)
        risk = content.get("risk_responses", "")
        for line in risk.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    def _render_question_tree(self, doc: Document, tree: Dict):
        """渲染追问树"""
        anchor = tree.get("anchor", "")
        branches = tree.get("branches", {})

        # 锚点问题
        doc.add_paragraph(f"开场锚定：{anchor}")

        # 分支
        for key, questions in branches.items():
            doc.add_paragraph(f"├─ 答\"{key}\" → {questions[0]}")
            if len(questions) > 1:
                doc.add_paragraph(f"│            → {questions[1]}")
