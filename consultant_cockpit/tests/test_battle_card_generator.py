# tests/test_battle_card_generator.py
"""作战卡生成器单元测试

根据设计文档 11.4.2 节测试用例：
- 双模式切换测试
- SKU召回测试
- Word生成测试
- 降级处理测试
"""
import pytest
from datetime import datetime
from unittest.mock import Mock, MagicMock
from io import BytesIO

from src.core.battle_card_generator import (
    BattleCard,
    BattleCardGenerator,
    InsufficientSkuError,
    PRESET_STRATEGY_TREE,
    PRESET_BUSINESS_TREE
)


class MockFeishuClient:
    """Mock 飞书客户端"""

    def get_client_profile(self, company: str):
        """返回模拟的客户档案"""
        if company == "高完整度客户":
            return {
                "record_id": "rec_001",
                "fields": {
                    "客户公司名": "高完整度客户",
                    "产品线": "储能系统/光伏设备",
                    "客户群体": "工商业用户",
                    "收入结构": "设备销售70%/运营服务30%",
                    "毛利结构": "设备15%/服务40%",
                    "交付情况": "已交付项目50+",
                    "资源分布": "华东60%/华南40%",
                    "战略目标": "从设备商转型运营商",
                    "显性诉求": "希望提升运营收入占比",
                    "隐性痛点": "设备销售毛利持续下降"
                }
            }
        elif company == "低完整度客户":
            return {
                "record_id": "rec_002",
                "fields": {
                    "客户公司名": "低完整度客户",
                    "产品线": "储能"
                }
            }
        else:
            return {"record_id": "rec_003", "fields": {}}

    def calc_completeness(self, profile):
        """计算完整度"""
        if not profile or not profile.get("fields"):
            return 0.0

        fields = profile.get("fields", {})
        filled_count = sum(1 for v in fields.values() if v and len(str(v)) >= 5)
        total_fields = 10

        return filled_count / total_fields

    def list_records(self):
        """返回记录列表"""
        return [
            self.get_client_profile("高完整度客户"),
            self.get_client_profile("低完整度客户")
        ]


class MockLLMClient:
    """Mock LLM客户端"""

    def generate(self, prompt: str, temperature: float = 0.7, max_tokens: int = 2000) -> str:
        """返回模拟的LLM响应"""
        if "诊断假设" in prompt:
            return "基于储能系统集成商商业模式案例，客户的核心问题很可能是设备销售毛利下滑，需要加速向运营服务转型。"
        elif "诊断问题" in prompt:
            return "1. 您现在的核心业务是什么？\n2. 未来3年的战略目标是什么？\n3. 目前最大的挑战是什么？"
        else:
            return "模拟生成的文本内容"


class MockKnowledgeRetriever:
    """Mock 知识召回器"""

    def recall_by_keywords(self, keywords, top_k=15):
        """返回模拟的SKU卡片"""
        from src.core.knowledge_retriever import SKUCard

        mock_skus = [
            SKUCard(id="sku_001", title="设备商转运营商路径", summary="从设备销售转向运营服务的典型案例", confidence="🟢", stage="商业模式"),
            SKUCard(id="sku_002", title="储能系统集成商商业模式", summary="工商业储能系统集成商的盈利模式分析", confidence="🟢", stage="商业模式"),
            SKUCard(id="sku_003", title="虚拟电厂聚合商案例", summary="负荷聚合商参与电力市场的路径", confidence="🟢", stage="战略梳理"),
            SKUCard(id="sku_004", title="分布式光伏运营模式", summary="分布式光伏的商业模式创新", confidence="🟡", stage="商业模式"),
            SKUCard(id="sku_005", title="重卡换电商业案例", summary="重卡换电站的盈利模型", confidence="🟡", stage="商业模式"),
            SKUCard(id="sku_006", title="微电网运营案例", summary="工业园区微电网运营实践", confidence="🟢", stage="战略梳理"),
            SKUCard(id="sku_007", title="综合能源服务转型", summary="传统能源企业转型综合能源服务", confidence="🟡", stage="战略梳理"),
            SKUCard(id="sku_008", title="储能电站运营", summary="独立储能电站的商业模式", confidence="🟢", stage="商业模式"),
            SKUCard(id="sku_009", title="电力交易代理", summary="电力市场化交易代理服务", confidence="🟢", stage="战略梳理"),
            SKUCard(id="sku_010", title="需求响应聚合", summary="需求响应资源聚合模式", confidence="🟡", stage="商业模式"),
        ]
        return mock_skus[:top_k]


# ============= 测试用例 =============

def test_battle_card_dataclass():
    """测试 BattleCard 数据结构"""
    card = BattleCard(
        company="测试公司",
        date="2026-05-02",
        consultant="张顾问",
        mode="hypothesis",
        completeness=0.8,
        content={"diagnosis_hypothesis": "测试假设"}
    )
    assert card.company == "测试公司"
    assert card.mode == "hypothesis"
    assert card.completeness == 0.8


def test_generator_mode_switch_high_completeness():
    """测试高完整度客户切换到验证假设版"""
    generator = BattleCardGenerator(
        feishu_client=MockFeishuClient(),
        llm_client=MockLLMClient(),
        knowledge_retriever=MockKnowledgeRetriever()
    )

    card = generator.generate("高完整度客户", "张顾问")

    assert card.mode == "hypothesis"
    assert card.completeness >= 0.6
    assert "diagnosis_hypothesis" in card.content
    assert "strategy_questions" in card.content


def test_generator_mode_switch_low_completeness():
    """测试低完整度客户切换到信息建立版"""
    generator = BattleCardGenerator(
        feishu_client=MockFeishuClient(),
        llm_client=MockLLMClient(),
        knowledge_retriever=MockKnowledgeRetriever()
    )

    card = generator.generate("低完整度客户", "张顾问")

    assert card.mode == "info_building"
    assert card.completeness < 0.6
    assert "missing_fields" in card.content
    assert "strategy_tree" in card.content


def test_generator_sku_recall():
    """测试SKU召回"""
    generator = BattleCardGenerator(
        feishu_client=MockFeishuClient(),
        llm_client=MockLLMClient(),
        knowledge_retriever=MockKnowledgeRetriever()
    )

    profile = generator.feishu_client.get_client_profile("高完整度客户")
    skus = generator._get_top_skus(profile, top_n=10)

    assert len(skus) >= 6  # MIN_SKU_COUNT
    assert all("id" in sku for sku in skus)
    assert all("confidence" in sku for sku in skus)


def test_generator_sku_ranking():
    """测试SKU优先级排序（按可信度）"""
    generator = BattleCardGenerator(
        feishu_client=MockFeishuClient(),
        llm_client=MockLLMClient(),
        knowledge_retriever=MockKnowledgeRetriever()
    )

    profile = generator.feishu_client.get_client_profile("高完整度客户")
    candidate_skus = generator._get_top_skus(profile)
    ranked_skus = generator._rank_skus(candidate_skus, profile)

    # 🟢 应排在前面
    green_skus = [s for s in ranked_skus if s["confidence"] == "🟢"]
    yellow_skus = [s for s in ranked_skus if s["confidence"] == "🟡"]

    # 验证排序：高可信度在前
    if green_skus and yellow_skus:
        first_green_idx = ranked_skus.index(green_skus[0])
        first_yellow_idx = ranked_skus.index(yellow_skus[0])
        assert first_green_idx <= first_yellow_idx


def test_generator_insufficient_sku_error():
    """测试SKU数量不足异常"""
    generator = BattleCardGenerator(
        feishu_client=MockFeishuClient(),
        llm_client=MockLLMClient(),
        knowledge_retriever=None  # 无知识召回器
    )

    # Mock 返回少量SKU
    generator._get_mock_skus = lambda n: [
        {"id": "sku_001", "title": "测试SKU", "confidence": "🔴", "stage": "测试"}
    ]

    profile = generator.feishu_client.get_client_profile("高完整度客户")

    with pytest.raises(InsufficientSkuError):
        generator._generate_hypothesis_version(profile)


def test_generator_word_output():
    """测试Word文档生成"""
    generator = BattleCardGenerator(
        feishu_client=MockFeishuClient(),
        llm_client=MockLLMClient(),
        knowledge_retriever=MockKnowledgeRetriever()
    )

    card = generator.generate("高完整度客户", "张顾问")
    word_bytes = generator.render_to_word(card)

    assert isinstance(word_bytes, bytes)
    assert len(word_bytes) > 0

    # 验证可以解析为Word文档
    from docx import Document
    from io import BytesIO
    doc = Document(BytesIO(word_bytes))
    assert len(doc.paragraphs) > 0


def test_generator_word_output_info_building():
    """测试信息建立版Word文档生成"""
    generator = BattleCardGenerator(
        feishu_client=MockFeishuClient(),
        llm_client=MockLLMClient(),
        knowledge_retriever=MockKnowledgeRetriever()
    )

    card = generator.generate("低完整度客户", "张顾问")
    word_bytes = generator.render_to_word(card)

    assert isinstance(word_bytes, bytes)

    from docx import Document
    from io import BytesIO
    doc = Document(BytesIO(word_bytes))

    # 验证包含追问树
    text_content = "\n".join([p.text for p in doc.paragraphs])
    assert "追问树" in text_content or "开场锚定" in text_content


def test_generator_llm_fallback():
    """测试LLM调用失败的降级处理"""
    # Mock LLM 会抛出异常
    failing_llm = Mock()
    failing_llm.generate = Mock(side_effect=Exception("LLM调用失败"))

    generator = BattleCardGenerator(
        feishu_client=MockFeishuClient(),
        llm_client=failing_llm,
        knowledge_retriever=MockKnowledgeRetriever()
    )

    profile = generator.feishu_client.get_client_profile("高完整度客户")
    skus = generator._get_top_skus(profile, top_n=3)

    # 调用 _generate_hypothesis，应该降级返回模板填充
    hypothesis = generator._generate_hypothesis(skus, profile)

    # 降级后应该有内容
    assert len(hypothesis) > 0
    assert "案例" in hypothesis or "经验" in hypothesis


def test_preset_question_trees():
    """测试预设追问树结构"""
    assert "anchor" in PRESET_STRATEGY_TREE
    assert "branches" in PRESET_STRATEGY_TREE
    assert len(PRESET_STRATEGY_TREE["branches"]) >= 3

    assert "anchor" in PRESET_BUSINESS_TREE
    assert "branches" in PRESET_BUSINESS_TREE
    assert len(PRESET_BUSINESS_TREE["branches"]) >= 3


def test_missing_fields_identification():
    """测试缺失字段识别"""
    generator = BattleCardGenerator(
        feishu_client=MockFeishuClient(),
        llm_client=MockLLMClient(),
        knowledge_retriever=None
    )

    # 低完整度客户
    profile = generator.feishu_client.get_client_profile("低完整度客户")
    missing = generator._identify_missing_fields(profile)

    assert len(missing) > 0
    assert "客户群体" in missing or "收入结构" in missing


def test_constraint_validation():
    """测试硬约束校验"""
    generator = BattleCardGenerator(
        feishu_client=MockFeishuClient(),
        llm_client=MockLLMClient(),
        knowledge_retriever=None
    )

    # 有🟢/🟡SKU的情况
    valid_skus = [
        {"id": "sku_001", "confidence": "🟢"},
        {"id": "sku_002", "confidence": "🟡"}
    ]
    content = {"diagnosis_hypothesis": "测试"}
    assert generator._validate_constraints(content, valid_skus) == True

    # 只有🔴SKU的情况
    invalid_skus = [
        {"id": "sku_001", "confidence": "🔴"}
    ]
    with pytest.raises(InsufficientSkuError):
        generator._validate_constraints(content, invalid_skus)